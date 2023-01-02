from docx import Document
import datetime
import openpyxl
import warnings

warnings.filterwarnings("ignore", category=UserWarning, module="openpuxl")

# IS EXCELIO
imone = input('Iveskite imones pavadinima:')
paketo_id = input('Iveskite projekto ID:')
skyriaus_pavadinimas = input('Iveskite skyriaus pavadinima:')
path1 = f"C:/Users/User/Dalgasgroup A S/HD Forest Lithuania - Documents/Evaluation/04_Pirkta 2022/{imone}/{paketo_id}/{paketo_id}.xlsx"

wb = openpyxl.load_workbook(path1)

sheet = wb['Sheet1']

kadastras = sheet['C5'].value
Adresas = sheet['D5'].value
plotas = sheet['F5'].value

# I WORDA KELIMAS
document = Document(
    f'C:/Users/User/Dalgasgroup A S/HD Forest Lithuania - Documents/NZT leidimai pirkti sklypus/Misku ukio paskirtis/{imone}/Pras_mu_sablonas.docx')
today = datetime.date.today()

table = document.tables[0]

for paragraph in document.paragraphs:
    if 'Skyriaus_pavadinimas' in paragraph.text:
        print(paragraph.text)
    if 'Skyriaus_pavadinimas' in paragraph.text:
        paragraph.text = paragraph.text.replace('Skyriaus_pavadinimas', f"{skyriaus_pavadinimas}")
    if 'Siandienos_data' in paragraph.text:
        paragraph.text = paragraph.text.replace('Siandienos_data', f'{today}')

for row in table.rows:
    for cell in row.cells:
        if 'Sklypo_adresas' in cell.text:
            cell.text = cell.text.replace('Sklypo_adresas', f'{Adresas}')
    for cell in row.cells:
        if cell.text == 'Kadastras':
            cell.text = kadastras
    for cell in row.cells:
        if '_plotas' in cell.text:
            cell.text = cell.text.replace('_plotas', f'{plotas}00')

uzvadinimas = input('Uzvadinimas:')

document.save(
    f'C:/Users/User/Dalgasgroup A S/HD Forest Lithuania - Documents/NZT leidimai pirkti sklypus/Misku ukio paskirtis/{imone}/{uzvadinimas}.docx')
