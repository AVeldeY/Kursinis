from docx import Document
import datetime
import warnings
import pandas as pd

warnings.filterwarnings("ignore", category=UserWarning, module="openpuxl")

# IS EXCELIO
imone = input('Iveskite imones pavadinima:')
paketo_id = input('Iveskite projekto ID:')
skyriaus_pavadinimas = input('Iveskite skyriaus pavadinima:')

path1 = 'C:/Users/User/Desktop/trylt-248.xlsx'
df = pd.read_excel(path1)

kadastras = df.iloc[4, 2]
Adresas = df.iloc[4, 3]
plotas = df.iloc[4, 5]

# I WORDA KELIMAS
document = Document(
    f'C:/Users/User/Dalgasgroup A S/HD Forest Lithuania - Documents/NZT leidimai pirkti sklypus/Misku ukio paskirtis/{imone}/Pras_mu_sablonas.docx')
today = datetime.date.today()

table = document.tables[0]

for paragraph in document.paragraphs:
    if 'Skyriaus_pavadinimas' in paragraph.text:
        print(paragraph.text)
    if 'Skyriaus_pavadinimas' in paragraph.text:
        paragraph.text = paragraph.text.replace('Skyriaus_pavadinimas', f"{skyriaus_pavadinimas}")
    if 'Siandienos_data' in paragraph.text:
        paragraph.text = paragraph.text.replace('Siandienos_data', f'{today}')

for row in table.rows:
    for cell in row.cells:
        if 'Sklypo_adresas' in cell.text:
            cell.text = cell.text.replace('Sklypo_adresas', f'{Adresas}')
    for cell in row.cells:
        if cell.text == 'Kadastras':
            cell.text = kadastras
    for cell in row.cells:
        if '_plotas' in cell.text:
            cell.text = cell.text.replace('_plotas', f'{plotas}')

uzvadinimas = input('Uzvadinimas:')

document.save(
    f'C:/Users/User/Dalgasgroup A S/HD Forest Lithuania - Documents/NZT leidimai pirkti sklypus/Misku ukio paskirtis/{imone}/{uzvadinimas}.docx')
